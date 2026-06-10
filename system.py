
import streamlit as st
from google import genai
import pandas as pd
import json
import re
from docx import Document
from pypdf import PdfReader
from io import BytesIO

# =========================
# CONFIG
# =========================

MODEL_NAME = "gemini-2.5-flash-lite"

client = genai.Client(
    api_key=st.secrets["GEMINI_API_KEY"]
)

# =========================
# UI
# =========================

st.set_page_config(
    page_title="MakeUrQuiz",
    page_icon="⚡",
    layout="wide"
)

st.title("⚡ MakeUrQuiz - Tạo bộ câu hỏi từ tài liệu học tập")
st.caption("Made by Đăng Khoa 🔰")


# =========================
# INPUT DATA
# =========================

st.header("📚 Nguồn dữ liệu")

col1, col2 = st.columns([2, 1])

with col1:

    source_type = st.selectbox(
        "Chọn nguồn dữ liệu",
        [
            "Dán văn bản",
            "Upload file"
        ]
    )

    raw_text = ""

    if source_type == "Dán văn bản":
        raw_text = st.text_area(
            "Dán nội dung tại đây",
            height=300
        )

        MAX_CHARS = 10000
        if len(raw_text) > MAX_CHARS:
            st.warning("Dữ liệu quá dài, chỉ lấy 10.000 ký tự đầu.")
            raw_text = raw_text[:MAX_CHARS]
        
    else:
        uploaded_file = st.file_uploader(
            "Upload file",
            type=["txt", "docx","pdf(chỉ định dạng kí tự)"]
        )

        if uploaded_file:

            try:

                if uploaded_file.name.endswith(".docx"):

                    doc = Document(uploaded_file)

                    raw_text = "\n".join(
                        para.text
                        for para in doc.paragraphs
                    )
                elif uploaded_file.name.endswith(".pdf"):

                    reader = PdfReader(uploaded_file)

                    raw_text = ""

                    for page in reader.pages:
                        text = page.extract_text()
                        if text:
                            raw_text += text + "\n"
                else:
                    raw_text = uploaded_file.read().decode("utf-8")

                MAX_CHARS = 10000

                if len(raw_text) > MAX_CHARS:
                    st.warning(
                        f"File quá dài ({len(raw_text):,} ký tự). "
                        f"Chỉ lấy {MAX_CHARS:,} ký tự đầu."
                    )
                    raw_text = raw_text[:MAX_CHARS]

            except Exception as e:
                st.error(f"Lỗi đọc file: {e}")
st.info(f"Đã đọc {len(raw_text):,} ký tự")

with col2:

    q_type = st.selectbox(
        "Loại câu hỏi",
        [
            "Trắc nghiệm",
            "Tự luận",
            "True/False"
        ]
    )

    num_q = st.number_input(
        "Số lượng câu hỏi",
        min_value=1,
        max_value=40,
        value=1
    )
    max_questions = max(1, len(raw_text) // 100)

    if num_q > max_questions:

        st.info(
            f"Tài liệu hơi ngắn. "
            f"Tự động giảm từ {num_q} xuống {max_questions} câu."
        )

        num_q = max_questions
    
    difficulty = st.selectbox(
        "Độ khó",
        [
            "Dễ",
            "Trung bình",
            "Khó"
        ]
    )

    generate_btn = st.button(
        "🚀 Tạo câu hỏi",
        use_container_width=True
    )


# =========================
# PROMPT
# =========================

def make_prompt(text, q_type, num_q, difficulty):

    if q_type == "Trắc nghiệm":

        instruction = f"""
Tạo {num_q} câu hỏi trắc nghiệm.

Độ khó: {difficulty}

Mỗi câu phải có:

- question
- options
- answer
- explanation

options phải gồm đúng 4 lựa chọn:
A, B, C, D

answer chỉ được là:
"A"
hoặc
"B"
hoặc
"C"
hoặc
"D"

KHÔNG được ghi nội dung đáp án.

Ví dụ:

{{
  "question": "...",
  "options": {{
    "A": "...",
    "B": "...",
    "C": "...",
    "D": "..."
  }},
  "answer": "C",
  "explanation": "..."
}}

Trả về JSON LIST duy nhất.
"""

    elif q_type == "Tự luận":

        instruction = f"""
Tạo {num_q} câu hỏi tự luận.

Độ khó: {difficulty}

Tạo {num_q} câu hỏi tự luận.

Mỗi câu phải có:

- question
- hint
- keywords

keywords là danh sách các từ khóa quan trọng.

Ví dụ:

{{
  "question": "...",
  "hint": "...",
  "keywords": ["...", "..."]
}}

Trả về JSON LIST duy nhất.
"""

    else:

        instruction = f"""
Tạo {num_q} câu hỏi đúng/sai.

Độ khó: {difficulty}

Mỗi câu gồm:
- câu hỏi
- đáp án (ONLY: TRUE/FALSE)
- giải thích

Trả về JSON LIST duy nhất.
"""
    return f"""
{instruction}

NỘI DUNG:

{text}

CHỈ TRẢ VỀ JSON.
KHÔNG markdown.
KHÔNG giải thích thêm.
"""

# =========================
# GENERATE
# =========================

if generate_btn:

    if len(raw_text.strip()) < 20:
        st.warning("Vui lòng nhập ít nhất 20 ký tự.")
        st.stop()

    with st.spinner("Đang tạo câu hỏi..."):
        try:
            prompt = make_prompt(
                raw_text,
                q_type,
                int(num_q),
                difficulty
            )

            response = client.models.generate_content(
                model=MODEL_NAME,
                contents=prompt
            )

            model_text = (response.text or "").strip()

            # Xử lý nếu model trả markdown
            model_text = re.sub(r"^```json\s*", "", model_text)
            model_text = re.sub(r"```$", "", model_text).strip()

            q_list = json.loads(model_text)

            # ====== Tạo bảng preview ======
            rows = []
            for idx, item in enumerate(q_list, start=1):
                options = item.get("options", {})

                if isinstance(options, dict):
                    a = options.get("A", "")
                    b = options.get("B", "")
                    c = options.get("C", "")
                    d = options.get("D", "")
                elif isinstance(options, list):
                    a = options[0] if len(options) > 0 else ""
                    b = options[1] if len(options) > 1 else ""
                    c = options[2] if len(options) > 2 else ""
                    d = options[3] if len(options) > 3 else ""
                else:
                    a = b = c = d = ""

                rows.append({
                    "id": idx,
                    "type": item.get("type", q_type),
                    "question": item.get("question", ""),
                    "A": a,
                    "B": b,
                    "C": c,
                    "D": d,
                    "answer": item.get("answer", ""),
                    "explanation": item.get("explanation", item.get("hint", "")),
                    "keywords": ", ".join(item.get("keywords", [])) if isinstance(item.get("keywords"), list) else ""
                })

            df_out = pd.DataFrame(rows)

            # ====== Tạo TXT ======
            txt_content = ""

            for i, item in enumerate(q_list, start=1):
                txt_content += f"Câu {i}: {item.get('question', '')}\n"

                options = item.get("options")

                if q_type == "Trắc nghiệm" and options:
                    if isinstance(options, dict):
                        txt_content += f"A. {options.get('A', '')}\n"
                        txt_content += f"B. {options.get('B', '')}\n"
                        txt_content += f"C. {options.get('C', '')}\n"
                        txt_content += f"D. {options.get('D', '')}\n"
                    elif isinstance(options, list):
                        labels = ["A", "B", "C", "D"]
                        for j, opt in enumerate(options[:4]):
                            txt_content += f"{labels[j]}. {opt}\n"

                if item.get("answer"):
                    txt_content += f"Đáp án: {item.get('answer')}\n"

                if item.get("explanation"):
                    txt_content += f"Giải thích: {item.get('explanation')}\n"

                if q_type == "Tự luận" and item.get("hint"):
                    txt_content += f"Gợi ý: {item.get('hint')}\n"

                if q_type == "Tự luận" and item.get("keywords"):
                    if isinstance(item["keywords"], list):
                        txt_content += "Từ khóa: " + ", ".join(item["keywords"]) + "\n"

                txt_content += "\n" + "=" * 50 + "\n\n"

            # ====== Tạo Word (.docx) ======

            from docx import Document

            doc = Document()
            doc.add_heading("MakeUrQuiz", level=1)

            for line in txt_content.split("\n"):
                doc.add_paragraph(line)

            doc_buffer = BytesIO()
            doc.save(doc_buffer)
            doc_buffer.seek(0)
            
            st.session_state["df_out"] = df_out
            st.session_state["txt_content"] = txt_content
            st.session_state["doc_buffer"] = doc_buffer.getvalue()
            st.session_state["q_list"] = q_list
            
        except Exception as e:
            st.error(f"Lỗi: {e}")    
if "df_out" in st.session_state:
    st.success(
        f"Đã tạo {len(st.session_state['df_out'])} câu hỏi."
    )

    st.subheader("📋 Preview")

    st.dataframe(
        st.session_state["df_out"],
        use_container_width=True
    )
    
            # ====== Download ======
if "txt_content" in st.session_state:
    st.download_button(
        "⬇️ Tải TXT",
        data=st.session_state["txt_content"],
        file_name="questions.txt",
        mime="text/plain"
    )

    st.download_button(
        "⬇️ Tải Word (.docx)",
        data=st.session_state["doc_buffer"],
        file_name="questions.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
